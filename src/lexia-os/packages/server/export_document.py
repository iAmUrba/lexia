import sys
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def generate(json_file, template_file, output_file):
    doc = Document(template_file)
    with open(json_file, 'r', encoding='utf-8') as f:
        paragraphs = json.load(f)
        
    for p_data in paragraphs:
        p = doc.add_paragraph()
        
        # Alineación
        align = p_data.get('align')
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        runs = p_data.get('runs', [])
        for i, run_data in enumerate(runs):
            if run_data.get('type') == 'image':
                # Skip images
                continue
                
            text = run_data.get('text', '')
            run = p.add_run(text)
            
            if run_data.get('bold'):
                run.bold = True
            if run_data.get('italic'):
                run.italic = True
                
            run.font.name = 'Century Gothic'
            run.font.size = Pt(11)
            
    doc.save(output_file)

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--data', required=True)
    parser.add_argument('--template', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()
    
    generate(args.data, args.template, args.output)
