import sys
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
                
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = 'Century Gothic'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def generate_docx(events_json_path, template_path, output_path, month_year_str, start_date_str):
    with open(events_json_path, 'r', encoding='utf-8') as f:
        events = json.load(f)
        
    doc = Document(template_path)
    
    # Create the table
    table = doc.add_table(rows=0, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Row 1: Month/Year header
    row0 = table.add_row()
    cell0 = row0.cells[0]
    cell0.merge(row0.cells[4])
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(month_year_str)
    set_run_font(run, size=11, bold=True)
    
    start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
    days_names = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES"]
    
    # Row 2: Days
    row1 = table.add_row()
    for i in range(5):
        cell = row1.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        day_date = start_date.day + i
        run1 = p.add_run(f"{days_names[i]}\n")
        set_run_font(run1, size=10, bold=True)
        run2 = p.add_run(str(day_date))
        set_run_font(run2, size=10, bold=True)
        
    # Group events by Day and by Slot
    slots = [
        {"label": "08:30", "min": 0, "max": 9},
        {"label": "10:30", "min": 10, "max": 12},
        {"label": "02:00", "min": 13, "max": 23},
    ]
    
    border_kwargs = {
        "top": {"sz": 4, "val": "single", "color": "000000"},
        "bottom": {"sz": 4, "val": "single", "color": "000000"},
        "start": {"sz": 4, "val": "single", "color": "000000"},
        "end": {"sz": 4, "val": "single", "color": "000000"},
    }

    for r in [row0, row1]:
        for c in r.cells:
            set_cell_border(c, **border_kwargs)
            
    RED = RGBColor(255, 0, 0)
    BLACK = RGBColor(0, 0, 0)
    GREEN = RGBColor(0, 176, 80)
            
    for slot in slots:
        row = table.add_row()
        for i in range(5):
            cell = row.cells[i]
            set_cell_border(cell, **border_kwargs)
            
            # Use timedelta so it wraps across months correctly
            target_date = start_date + timedelta(days=i)
            
            slot_events = []
            for e in events:
                try:
                    dt = datetime.fromisoformat(e['datetime'])
                    if dt.year == target_date.year and dt.month == target_date.month and dt.day == target_date.day:
                        if slot["min"] <= dt.hour <= slot["max"]:
                            slot_events.append(e)
                except ValueError:
                    continue
                    
            if not slot_events:
                continue
                
            p = cell.paragraphs[0]
            
            for idx, ev in enumerate(slot_events):
                if idx > 0:
                    p = cell.add_paragraph()
                    
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                dt = datetime.fromisoformat(ev['datetime'])
                hh = dt.hour
                mm = dt.minute
                if hh > 12:
                    hh -= 12
                if hh == 0:
                    hh = 12
                time_str = f"{hh}:{mm:02d}\n"
                
                run_time = p.add_run(time_str)
                set_run_font(run_time, size=10, bold=True, color=RED)
                
                desc = ev.get('description', '')
                if '-' in desc:
                    parts = desc.split('-', 1)
                    tipo = parts[0].strip() + "\n"
                    detalle = parts[1].strip()
                else:
                    words = desc.split(' ')
                    tipo = " ".join(words[:2]) + "\n"
                    detalle = " ".join(words[2:])
                    
                run_tipo = p.add_run(tipo)
                set_run_font(run_tipo, size=10, bold=True, color=RED)
                
                is_notified = '✔' in ev.get('title', '') or ev.get('title') == 'Notificado'
                
                run_det = p.add_run(detalle)
                set_run_font(run_det, size=10, bold=True, color=BLACK)
                
                if is_notified:
                    run_check = p.add_run(" ✔")
                    set_run_font(run_check, size=10, bold=True, color=GREEN)
                
                if idx < len(slot_events) - 1:
                    p.add_run("\n") # Just add some spacing if there's another event
                    
    doc.save(output_path)

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--events', required=True)
    parser.add_argument('--template', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--month_str', required=True)
    parser.add_argument('--start_date', required=True)
    
    args = parser.parse_args()
    generate_docx(args.events, args.template, args.output, args.month_str, args.start_date)
